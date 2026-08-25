#!/usr/bin/env python3
"""Build the branded Capoeira Luanda songbook as DOCX and PDF.

The source repository stores each song, quadra, or rhythm note as an
extensionless UTF-8 file. This builder preserves the source wording and order
within each song and quadra, removes only the structural title line from rhythm
notes, skips one duplicated rhythm file, and publishes usable source/listen
links.
"""

from __future__ import annotations

import hashlib
import html
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    Image,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
)
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "assets" / "capoeira-luanda-blue-logo.png"
DOCX_OUT = ROOT / "output" / "docx" / "Capoeira_Luanda_Songbook_Blue_Edition.docx"
PDF_OUT = ROOT / "output" / "pdf" / "Capoeira_Luanda_Songbook_Blue_Edition.pdf"

PRIMARY = "0566AB"
DEEP = "133A56"
SKY = "5BA0CA"
PALE = "EAF4FA"
INK = "202020"
MUTED = "5A6670"
WHITE = "FFFFFF"

URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)
UNAVAILABLE_SOURCE_DOMAINS = {"capoeira-music.net"}


@dataclass
class Entry:
    title: str
    category: str
    lines: list[str]
    urls: list[str]
    bilingual: bool
    source_path: Path


def sort_key(value: str) -> str:
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", value).casefold().strip()


def url_domain(url: str) -> str:
    return urlparse(url).netloc.casefold().removeprefix("www.")


def is_publishable_url(url: str) -> bool:
    return url_domain(url) not in UNAVAILABLE_SOURCE_DOMAINS


def load_entries() -> list[Entry]:
    excluded_parts = {".git", "assets", "output", "tmp", "tools", "Mestres de Capoeira"}
    candidates: list[Path] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        if any(part in excluded_parts for part in rel.parts):
            continue
        if any(part.startswith(".") for part in rel.parts):
            continue
        if path.name == "README.md" or path.suffix.lower() in {".mov", ".jpg", ".jpeg", ".png"}:
            continue
        candidates.append(path)

    entries: list[Entry] = []
    seen_hashes: set[str] = set()
    for path in sorted(candidates, key=lambda item: sort_key(str(item.relative_to(ROOT)))):
        raw = path.read_bytes()
        digest = hashlib.sha1(raw).hexdigest()
        if digest in seen_hashes:
            continue
        seen_hashes.add(digest)

        text = unicodedata.normalize("NFC", raw.decode("utf-8-sig")).replace("\r\n", "\n").replace("\r", "\n")
        source_lines = [line.rstrip() for line in text.split("\n")]
        while source_lines and not source_lines[-1].strip():
            source_lines.pop()
        if not source_lines:
            continue

        rel_parts = path.relative_to(ROOT).parts
        title = source_lines[0].strip()
        if "Cuadras de Bimba" in rel_parts:
            category = "Quadras de Bimba"
        elif "Rythms" in rel_parts or title.casefold().endswith(" rhythm"):
            category = "Rhythms"
        else:
            category = "Songs"

        # In songs and quadras, the first source line is both the entry title
        # and the opening lyric. Rhythm-note files use it only as a structural
        # title, so only those entries begin their body on the following line.
        body = source_lines[1:] if category == "Rhythms" else source_lines[:]
        while body and not body[0].strip():
            body.pop(0)

        urls: list[str] = []
        content_lines: list[str] = []
        for line in body:
            if URL_RE.match(line.strip()):
                url = line.strip()
                if is_publishable_url(url):
                    urls.append(url)
            else:
                content_lines.append(line)
        while content_lines and not content_lines[-1].strip():
            content_lines.pop()

        bilingual = "(EN)" in title.upper() or any(line.strip().casefold() == "english" for line in content_lines)
        entries.append(
            Entry(
                title=title,
                category=category,
                lines=content_lines,
                urls=urls,
                bilingual=bilingual,
                source_path=path.relative_to(ROOT),
            )
        )

    category_order = {"Songs": 0, "Quadras de Bimba": 1, "Rhythms": 2}
    entries.sort(key=lambda entry: (category_order[entry.category], sort_key(entry.title)))
    return entries


def stanzas(lines: list[str]) -> list[list[str]]:
    groups: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if line.strip():
            current.append(line.strip())
        elif current:
            groups.append(current)
            current = []
    if current:
        groups.append(current)
    return groups


def register_pdf_fonts() -> None:
    font_dir = Path("/System/Library/Fonts/Supplemental")
    pdfmetrics.registerFont(TTFont("Arial", str(font_dir / "Arial.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Bold", str(font_dir / "Arial Bold.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Italic", str(font_dir / "Arial Italic.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-BoldItalic", str(font_dir / "Arial Bold Italic.ttf")))
    pdfmetrics.registerFontFamily(
        "Arial",
        normal="Arial",
        bold="Arial-Bold",
        italic="Arial-Italic",
        boldItalic="Arial-BoldItalic",
    )


class SectionBand(Flowable):
    def __init__(self, number: str, title: str, subtitle: str, width: float):
        super().__init__()
        self.number = number
        self.title = title
        self.subtitle = subtitle
        self.width = width
        self.height = 4.7 * inch

    def draw(self):
        c = self.canv
        c.saveState()
        c.setFillColor(HexColor(f"#{PALE}"))
        c.roundRect(0, 0.35 * inch, self.width, 3.85 * inch, 18, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{PRIMARY}"))
        c.circle(0.76 * inch, 3.44 * inch, 0.46 * inch, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Arial-Bold", 21)
        c.drawCentredString(0.76 * inch, 3.33 * inch, self.number)
        c.setFillColor(HexColor(f"#{DEEP}"))
        c.setFont("Arial-Bold", 29)
        c.drawString(0.45 * inch, 2.42 * inch, self.title)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.setFont("Arial", 12.5)
        text = c.beginText(0.47 * inch, 2.02 * inch)
        text.setLeading(17)
        for line in self.subtitle.split("\n"):
            text.textLine(line)
        c.drawText(text)
        c.setStrokeColor(HexColor(f"#{SKY}"))
        c.setLineWidth(2)
        c.line(0.47 * inch, 1.39 * inch, self.width - 0.47 * inch, 1.39 * inch)
        c.restoreState()


class SongbookDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs):
        super().__init__(filename, **kwargs)
        self._bookmark_counter = 0

    def beforeDocument(self):
        # multiBuild can paginate more than once while resolving the TOC. Keep
        # bookmark IDs deterministic across passes so the index can converge.
        self._bookmark_counter = 0

    def afterFlowable(self, flowable):
        style_name = getattr(getattr(flowable, "style", None), "name", "")
        if style_name not in {"PDF Section Title", "PDF Song Title"}:
            return
        level = 0 if style_name == "PDF Section Title" else 1
        title = flowable.getPlainText()
        key = f"bookmark-{self._bookmark_counter}"
        self._bookmark_counter += 1
        self.canv.bookmarkPage(key)
        self.canv.addOutlineEntry(title, key, level=level, closed=level == 0)
        self.notify("TOCEntry", (level, title, self.page, key))


def cover_page(canvas, doc):
    width, height = LETTER
    canvas.saveState()
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, width, height, fill=1, stroke=0)
    canvas.setFillColor(HexColor(f"#{PRIMARY}"))
    canvas.rect(0, height - 0.42 * inch, width, 0.42 * inch, fill=1, stroke=0)
    canvas.rect(0, 0, width, 0.64 * inch, fill=1, stroke=0)
    canvas.setFillColor(HexColor(f"#{PALE}"))
    canvas.circle(width + 0.58 * inch, height * 0.54, 2.4 * inch, fill=1, stroke=0)
    canvas.setFillColor(HexColor(f"#{SKY}"))
    canvas.circle(-0.9 * inch, height * 0.16, 1.75 * inch, fill=1, stroke=0)
    canvas.restoreState()


def content_page(canvas, doc):
    width, height = LETTER
    canvas.saveState()
    canvas.setStrokeColor(HexColor(f"#{SKY}"))
    canvas.setLineWidth(0.7)
    canvas.line(0.82 * inch, height - 0.57 * inch, width - 0.82 * inch, height - 0.57 * inch)
    canvas.setFillColor(HexColor(f"#{MUTED}"))
    canvas.setFont("Arial-Bold", 7.8)
    canvas.drawString(0.82 * inch, height - 0.43 * inch, "CAPOEIRA LUANDA")
    canvas.setFont("Arial", 7.8)
    canvas.drawRightString(width - 0.82 * inch, height - 0.43 * inch, "BLUE EDITION  ·  SONG BOOK")
    canvas.setStrokeColor(HexColor("#D2E5F1"))
    canvas.line(0.82 * inch, 0.54 * inch, width - 0.82 * inch, 0.54 * inch)
    canvas.setFillColor(HexColor(f"#{MUTED}"))
    canvas.setFont("Arial", 7.6)
    canvas.drawString(0.82 * inch, 0.34 * inch, "Songs · Quadras · Rhythms")
    canvas.drawRightString(width - 0.82 * inch, 0.34 * inch, str(doc.page))
    canvas.restoreState()


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "cover_kicker": ParagraphStyle(
            "Cover Kicker", fontName="Arial-Bold", fontSize=9.5, leading=12,
            textColor=HexColor(f"#{PRIMARY}"), alignment=TA_CENTER, spaceAfter=8,
        ),
        "cover_title": ParagraphStyle(
            "Cover Title", fontName="Arial-Bold", fontSize=30, leading=33,
            textColor=HexColor(f"#{DEEP}"), alignment=TA_CENTER, spaceAfter=8,
        ),
        "cover_subtitle": ParagraphStyle(
            "Cover Subtitle", fontName="Arial", fontSize=13, leading=17,
            textColor=HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=14,
        ),
        "front_title": ParagraphStyle(
            "Front Title", fontName="Arial-Bold", fontSize=22, leading=26,
            textColor=HexColor(f"#{DEEP}"), spaceBefore=12, spaceAfter=14,
        ),
        "body": ParagraphStyle(
            "PDF Body", fontName="Arial", fontSize=10.5, leading=14.3,
            textColor=HexColor(f"#{INK}"), spaceAfter=7,
        ),
        "small": ParagraphStyle(
            "PDF Small", fontName="Arial", fontSize=8.5, leading=11.2,
            textColor=HexColor(f"#{MUTED}"), spaceAfter=5,
        ),
        "section_title": ParagraphStyle(
            "PDF Section Title", fontName="Arial-Bold", fontSize=1, leading=1,
            textColor=colors.white, spaceAfter=0,
        ),
        "song_title": ParagraphStyle(
            "PDF Song Title", fontName="Arial-Bold", fontSize=20, leading=23,
            textColor=HexColor(f"#{PRIMARY}"), spaceAfter=5,
        ),
        "song_meta": ParagraphStyle(
            "PDF Song Meta", fontName="Arial-Bold", fontSize=8.1, leading=10,
            textColor=HexColor(f"#{MUTED}"), spaceAfter=11,
        ),
        "stanza": ParagraphStyle(
            "PDF Stanza", fontName="Arial", fontSize=10.35, leading=13.8,
            textColor=HexColor(f"#{INK}"), spaceAfter=4.5,
            allowWidows=0, allowOrphans=0,
        ),
        "translation": ParagraphStyle(
            "PDF Translation Label", fontName="Arial-Bold", fontSize=8.5, leading=11,
            textColor=HexColor(f"#{PRIMARY}"), spaceBefore=5, spaceAfter=5,
        ),
        "source": ParagraphStyle(
            "PDF Source", fontName="Arial", fontSize=7.7, leading=10,
            textColor=HexColor(f"#{PRIMARY}"), spaceBefore=0, spaceAfter=0,
        ),
    }


def escape_lines(lines: list[str]) -> str:
    return "<br/>".join(html.escape(line) for line in lines)


def build_pdf(entries: list[Entry]) -> None:
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    register_pdf_fonts()
    styles = pdf_styles()
    page_width, page_height = LETTER
    frame = Frame(
        0.82 * inch,
        0.72 * inch,
        page_width - 1.64 * inch,
        page_height - 1.48 * inch,
        leftPadding=0,
        bottomPadding=0,
        rightPadding=0,
        topPadding=0,
        id="content-frame",
    )
    cover_frame = Frame(
        0.9 * inch,
        0.78 * inch,
        page_width - 1.8 * inch,
        page_height - 1.56 * inch,
        leftPadding=0,
        bottomPadding=0,
        rightPadding=0,
        topPadding=0,
        id="cover-frame",
    )
    doc = SongbookDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        leftMargin=0.82 * inch,
        rightMargin=0.82 * inch,
        topMargin=0.74 * inch,
        bottomMargin=0.72 * inch,
        title="Capoeira Luanda Song Book - Blue Edition",
        author="Capoeira Luanda",
        subject="Songs, Quadras de Bimba, and rhythm references",
        creator="Capoeira Luanda Song Book Builder",
    )
    doc.addPageTemplates(
        [
            PageTemplate(id="cover", frames=[cover_frame], onPage=cover_page),
            PageTemplate(id="content", frames=[frame], onPage=content_page),
        ]
    )

    story: list[Flowable] = []
    story.append(Spacer(1, 0.54 * inch))
    logo = Image(str(LOGO), width=3.15 * inch, height=3.15 * inch)
    logo.hAlign = "CENTER"
    story.append(logo)
    story.append(Spacer(1, 0.23 * inch))
    story.append(Paragraph("CAPOEIRA LUANDA", styles["cover_kicker"]))
    story.append(Paragraph("SONG BOOK", styles["cover_title"]))
    story.append(Paragraph("Songs · Quadras de Bimba · Rhythms", styles["cover_subtitle"]))
    story.append(Spacer(1, 0.44 * inch))
    story.append(Paragraph("BLUE EDITION", styles["cover_kicker"]))
    story.append(NextPageTemplate("content"))
    story.append(PageBreak())

    story.append(Paragraph("About this edition", styles["front_title"]))
    story.append(
        Paragraph(
            "This song book gathers the complete set of unique text entries in the local "
            "<b>LuandaSongBook</b> collection. Songs are followed by the Quadras de Bimba and "
            "rhythm reference notes.",
            styles["body"],
        )
    )
    story.append(
        Paragraph(
            "The wording, capitalization, translations, song-specific spellings, and opening lyric "
            "of every song are preserved from the source files. Unavailable legacy web addresses, "
            "rhythm-note headings, and one exact duplicate rhythm file were normalized for publication.",
            styles["body"],
        )
    )
    counts = {category: sum(1 for entry in entries if entry.category == category) for category in ("Songs", "Quadras de Bimba", "Rhythms")}
    story.append(Spacer(1, 0.12 * inch))
    story.append(
        Paragraph(
            f"<b>{counts['Songs']} songs</b> &nbsp;&nbsp;·&nbsp;&nbsp; "
            f"<b>{counts['Quadras de Bimba']} quadras</b> &nbsp;&nbsp;·&nbsp;&nbsp; "
            f"<b>{counts['Rhythms']} rhythm notes</b>",
            styles["body"],
        )
    )
    story.append(Spacer(1, 0.18 * inch))
    story.append(
        Paragraph(
            "Entries marked <b>PORTUGUÊS + ENGLISH</b> include an English translation in the "
            "source. Active source and listening links are retained at the end of each entry when available.",
            styles["small"],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Contents", styles["front_title"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle(
            "TOC Section", fontName="Arial-Bold", fontSize=10.5, leading=14,
            leftIndent=0, firstLineIndent=0, textColor=HexColor(f"#{DEEP}"), spaceBefore=8,
        ),
        ParagraphStyle(
            "TOC Entry", fontName="Arial", fontSize=8.9, leading=11.4,
            leftIndent=13, firstLineIndent=0, textColor=HexColor(f"#{INK}"),
        ),
    ]
    toc.dotsMinLevel = 0
    story.append(toc)

    category_info = [
        ("01", "Songs", "Cantigas for the roda, collected in one alphabetical section."),
        ("02", "Quadras de Bimba", "Traditional quadras associated with Mestre Bimba."),
        ("03", "Rhythms", "Short playing notes and listening references for key rhythms."),
    ]
    for number, category, subtitle in category_info:
        category_entries = [entry for entry in entries if entry.category == category]
        story.append(PageBreak())
        story.append(Paragraph(category, styles["section_title"]))
        story.append(SectionBand(number, category, subtitle, frame._width))
        story.append(PageBreak())
        for index, entry in enumerate(category_entries):
            if index:
                story.append(PageBreak())
            meta = category.upper()
            if entry.bilingual:
                meta += "  ·  PORTUGUÊS + ENGLISH"
            story.append(
                KeepTogether(
                    [
                        Paragraph(html.escape(entry.title), styles["song_title"]),
                        Paragraph(meta, styles["song_meta"]),
                    ]
                )
            )
            groups = stanzas(entry.lines)
            if not groups:
                story.append(Paragraph("No text supplied in the source file.", styles["small"]))
            for group in groups:
                normalized = " ".join(group).strip().casefold()
                if len(group) == 1 and normalized in {"english", "english translation", "português", "portuguese"}:
                    label = "ENGLISH TRANSLATION" if "english" in normalized else "PORTUGUÊS"
                    story.append(Paragraph(label, styles["translation"]))
                elif len(group) == 1 and normalized.startswith("includes english translation"):
                    story.append(Paragraph(html.escape(group[0]).upper(), styles["translation"]))
                else:
                    story.append(Paragraph(escape_lines(group), styles["stanza"]))
            for url in entry.urls:
                domain = urlparse(url).netloc.removeprefix("www.") or "source"
                safe_url = html.escape(url, quote=True)
                story.append(
                    Paragraph(
                        f'<link href="{safe_url}" color="#{PRIMARY}"><b>LISTEN / SOURCE</b> · {html.escape(domain)}</link>',
                        styles["source"],
                    )
                )

    doc.multiBuild(story)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Arial") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_style(style, size, color, bold, before, after, line_spacing) -> None:
    style.font.name = "Arial"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_outline_level(style, level: int) -> None:
    p_pr = style._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def paragraph_bottom_border(paragraph, color: str, size: int = 8, space: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=8, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PRIMARY)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(color)
    run_pr.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update this field to display the full contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    set_run_font(run, size=10, color=MUTED)


def set_docx_properties(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_repeat_style(normal, 11, INK, False, 0, 6, 1.25)
    set_repeat_style(doc.styles["Heading 1"], 16, PRIMARY, True, 18, 10, 1.0)
    set_repeat_style(doc.styles["Heading 2"], 13, PRIMARY, True, 14, 7, 1.0)
    set_repeat_style(doc.styles["Heading 3"], 12, DEEP, True, 10, 5, 1.0)

    subtitle = doc.styles["Subtitle"]
    set_repeat_style(subtitle, 13, MUTED, False, 0, 10, 1.15)

    styles = doc.styles
    song_title = styles.add_style("Song Title", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(song_title, 19, PRIMARY, True, 0, 5, 1.0)
    add_outline_level(song_title, 1)
    song_title.paragraph_format.keep_with_next = True

    section_title = styles.add_style("Section Title", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(section_title, 28, DEEP, True, 0, 8, 1.0)
    add_outline_level(section_title, 0)

    lyrics = styles.add_style("Lyrics", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(lyrics, 10.5, INK, False, 0, 7, 1.15)
    lyrics.paragraph_format.widow_control = True

    song_meta = styles.add_style("Song Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(song_meta, 8.3, MUTED, True, 0, 10, 1.0)
    song_meta.paragraph_format.keep_with_next = True

    translation = styles.add_style("Translation Label", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(translation, 8.5, PRIMARY, True, 5, 5, 1.0)

    source = styles.add_style("Source Link", WD_STYLE_TYPE.PARAGRAPH)
    set_repeat_style(source, 8, PRIMARY, False, 3, 2, 1.0)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = "Capoeira Luanda Song Book - Blue Edition"
    props.subject = "Songs, Quadras de Bimba, and rhythm references"
    props.author = "Capoeira Luanda"
    props.keywords = "Capoeira Luanda, song book, cantigas, quadras, rhythms"


def build_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    left = p.add_run("CAPOEIRA LUANDA")
    set_run_font(left, size=8, color=PRIMARY, bold=True)
    right = p.add_run("    BLUE EDITION  ·  SONG BOOK")
    set_run_font(right, size=8, color=MUTED)
    paragraph_bottom_border(p, SKY, size=6, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Songs · Quadras · Rhythms    ")
    set_run_font(run, size=8, color=MUTED)
    add_page_number(p)


def add_docx_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(3.25))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", "Capoeira Luanda blue logo")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("CAPOEIRA LUANDA")
    set_run_font(run, size=10, color=PRIMARY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("SONG BOOK")
    set_run_font(run, size=31, color=DEEP, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Songs · Quadras de Bimba · Rhythms")
    set_run_font(run, size=13, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BLUE EDITION")
    set_run_font(run, size=9, color=PRIMARY, bold=True)


def add_docx_section_page(doc: Document, number: str, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(number)
    set_run_font(run, size=19, color=WHITE, bold=True)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PRIMARY)
    p_pr.append(shd)

    p = doc.add_paragraph(title, style="Section Title")
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(subtitle)
    set_run_font(run, size=12.5, color=MUTED)
    paragraph_bottom_border(p, SKY, size=12, space=14)


def build_docx(entries: list[Entry]) -> None:
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_docx_properties(doc)
    build_header_footer(doc)
    add_docx_cover(doc)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("About this edition")
    set_run_font(run, size=22, color=DEEP, bold=True)
    doc.add_paragraph(
        "This song book gathers the complete set of unique text entries in the local LuandaSongBook collection. "
        "Songs are followed by the Quadras de Bimba and rhythm reference notes."
    )
    doc.add_paragraph(
        "The wording, capitalization, translations, song-specific spellings, and opening lyric of every song are "
        "preserved from the source files. Unavailable legacy web addresses, rhythm-note headings, and one exact "
        "duplicate rhythm file were normalized for publication."
    )
    counts = {category: sum(1 for entry in entries if entry.category == category) for category in ("Songs", "Quadras de Bimba", "Rhythms")}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    for index, (label, value) in enumerate(
        (("songs", counts["Songs"]), ("quadras", counts["Quadras de Bimba"]), ("rhythm notes", counts["Rhythms"]))
    ):
        if index:
            sep = p.add_run("   ·   ")
            set_run_font(sep, size=10.5, color=SKY)
        number = p.add_run(str(value))
        set_run_font(number, size=11, color=DEEP, bold=True)
        label_run = p.add_run(f" {label}")
        set_run_font(label_run, size=10.5, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Entries marked PORTUGUÊS + ENGLISH include an English translation in the source. "
        "Active source and listening links are retained at the end of each entry when available."
    )
    set_run_font(run, size=9, color=MUTED)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Contents")
    set_run_font(run, size=22, color=DEEP, bold=True)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(8)
    add_toc(toc_p)
    p = doc.add_paragraph()
    run = p.add_run("Word updates this table automatically when the document opens. If needed, choose Update Field.")
    set_run_font(run, size=8.5, color=MUTED, italic=True)

    category_info = [
        ("01", "Songs", "Cantigas for the roda, collected in one alphabetical section."),
        ("02", "Quadras de Bimba", "Traditional quadras associated with Mestre Bimba."),
        ("03", "Rhythms", "Short playing notes and listening references for key rhythms."),
    ]
    for number, category, subtitle in category_info:
        doc.add_page_break()
        add_docx_section_page(doc, number, category, subtitle)
        for entry in [item for item in entries if item.category == category]:
            p = doc.add_paragraph(entry.title, style="Song Title")
            # A paragraph containing a hard page-break can itself roll onto the
            # next page when the preceding song is full, producing a blank page.
            # Page-break-before on the title is deterministic across Word/Pages.
            p.paragraph_format.page_break_before = True
            p.paragraph_format.keep_with_next = True
            meta = category.upper()
            if entry.bilingual:
                meta += "  ·  PORTUGUÊS + ENGLISH"
            doc.add_paragraph(meta, style="Song Meta")
            groups = stanzas(entry.lines)
            if not groups:
                p = doc.add_paragraph("No text supplied in the source file.")
                set_run_font(p.runs[0], size=9, color=MUTED, italic=True)
            for group in groups:
                normalized = " ".join(group).strip().casefold()
                if len(group) == 1 and normalized in {"english", "english translation", "português", "portuguese"}:
                    label = "ENGLISH TRANSLATION" if "english" in normalized else "PORTUGUÊS"
                    doc.add_paragraph(label, style="Translation Label")
                elif len(group) == 1 and normalized.startswith("includes english translation"):
                    doc.add_paragraph(group[0].upper(), style="Translation Label")
                else:
                    p = doc.add_paragraph(style="Lyrics")
                    for line_index, line in enumerate(group):
                        if line_index:
                            p.add_run().add_break()
                        run = p.add_run(line)
                        set_run_font(run, size=10.5, color=INK)
            for url in entry.urls:
                domain = urlparse(url).netloc.removeprefix("www.") or "source"
                p = doc.add_paragraph(style="Source Link")
                label = p.add_run("LISTEN / SOURCE · ")
                set_run_font(label, size=8, color=PRIMARY, bold=True)
                add_hyperlink(p, domain, url)

    doc.save(DOCX_OUT)


def main() -> None:
    if not LOGO.exists():
        raise FileNotFoundError(f"Logo not found: {LOGO}")
    entries = load_entries()
    if not entries:
        raise RuntimeError("No songbook text entries found")
    build_docx(entries)
    build_pdf(entries)
    counts = {category: sum(1 for entry in entries if entry.category == category) for category in ("Songs", "Quadras de Bimba", "Rhythms")}
    print(f"Built {DOCX_OUT}")
    print(f"Built {PDF_OUT}")
    print(f"Entries: {len(entries)} ({counts})")


if __name__ == "__main__":
    main()
